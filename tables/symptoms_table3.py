# Generate supplementary table 3: Questionnaires
# run in python 3.12 environment

import os
from pathlib import Path
import pandas as pd

import matplotlib.pyplot as plt
import matplotlib.colors as mcolors

from docx import Document

# setting for nature publishing
plt.rcParams['pdf.fonttype']=42

# linux doesn't have Arial
plt.rcParams.update({
    "font.family": "sans-serif",
    "font.sans-serif": ["DejaVu Sans"],
    "font.size": 12,
    "axes.labelsize": 14,
    "axes.titlesize": 14,
    "xtick.labelsize": 12,
    "ytick.labelsize": 13,
    "legend.fontsize": 12,
})

# --------------------------------------------------
# Configuration
# --------------------------------------------------
home_dir = Path("L:/Lab_LucaC/Carina/")
csv_path = Path(f"{home_dir}/canonical_hmm_finalsample/clinical_demo_combined_012026.csv")
fig_dir = Path(f'{home_dir}/canonical_hmm_finalsample/figures')

if not os.path.exists(fig_dir):
    os.makedirs(fig_dir)

# Colorblind-friendly palette (Okabe-Ito)
cb_palette = ["#E69F00", "#56B4E9", "#009E73", "#F0E442", "#0072B2",
              "#D55E00", "#CC79A7", "#999999", "#000000"]

base_color = cb_palette[0]

def lighten_color(color, amount=0.6):
    r, g, b = mcolors.to_rgb(color)
    return (
        r + (1 - r) * amount,
        g + (1 - g) * amount,
        b + (1 - b) * amount,
    )

light_color = lighten_color(base_color, amount=0.6)

# --------------------------------------------------
# Data loading & preprocessing
# --------------------------------------------------
def load_and_prep_data(exclude_repeater: bool = False):

    df = pd.read_csv(csv_path)

    # ensure patient IDs are strings
    df["patient"] = df["patient"].astype(str)

    if exclude_repeater:
        df = df[~df["patient"].str.contains("R", na=False)]
    
    # load patient ID list
    patient_ids = pd.read_csv(Path("L:/Lab_LucaC/Carina/canonical_hmm_finalsample/prepared_data_giles_1Hz_3Hzfiltereddata/patients_fitted_for_this_hmm.csv"))

    patient_ids["patient_id"] = patient_ids["patient_id"].astype(str)

    # keep only patients fitted for this HMM
    df = df[df["patient"].isin(patient_ids["patient_id"])]

    print(f"Analyzing {df['patient'].nunique()} patients")

    # compute years with depression
    df["years_with_depression"] = df["age"] - df["age of symptom onset"]

    return df

df = load_and_prep_data()

# prep session column

df['session'] = df['session'].str.replace(" ", "_")

# create a table with all questionnaires (HADS-D, HAM-A, MADRS, HADS-A)
questionnaires = {
    "MADRS": {
        "col": "madrs_total",
        "timepoints": ["pre", "post"]
    },
    "HAMA": {
        "col": "hama_total",
        "timepoints": ["pre", "post"]
    },
    "HADSD": {
        "col": "hads_dep_total",
        "timepoints": ["pre", "week_1", "week_2", "week_3", "week_4", "post"]
    },
    "HADSA": {
        "col": "hads_anx_total",
        "timepoints": ["pre", "week_1", "week_2", "week_3", "week_4", "post"]
    }
}

timepoints = ["pre", "week_1", "week_2", "week_3", "week_4", "post"]
rows = []

for tp in timepoints:
    row = {"Timepoint": tp}

    for name, cfg in questionnaires.items():
        col = cfg["col"]

        if tp in cfg["timepoints"]:
            row[name] = median_iqr_n(
                df.loc[df["session"] == tp, col]
            )
        else:
            row[name] = "–"

    rows.append(row)

# Add change row
change_row = {"Timepoint": "Change baseline → post"}
for name, cfg in questionnaires.items():
    change_row[name] = change_baseline_to_post_minmax(df, cfg["col"])

rows.append(change_row)
table_symptoms = pd.DataFrame(rows)

save_dir = Path("C:/Users/CarinaF/tms_brainpattern/tables")

dataframe_to_word_table(
        table_symptoms,
        f"{save_dir}/Table3_Symptoms.docx",
        title="Table 3. Symptoms"
    )



def dataframe_to_word_table(df, filename, title=None):
    doc = Document()

    if title:
        p = doc.add_paragraph(title)
        p.style = "Normal"
        p.runs[0].bold = True

    table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1])
    table.style = "Table Grid"

    # Header
    for j, col in enumerate(df.columns):
        cell = table.cell(0, j)
        cell.text = col
        cell.paragraphs[0].runs[0].bold = True

    # Body
    for i in range(df.shape[0]):
        for j in range(df.shape[1]):
            table.cell(i + 1, j).text = str(df.iat[i, j])

    doc.save(filename)



def median_iqr_n(series, digits=1):
    s = pd.to_numeric(series, errors="coerce")
    n = s.notna().sum()
    if n == 0:
        return "–"
    s = s.dropna()
    q1, q3 = s.quantile([0.25, 0.75])
    return f"{s.median():.{digits}f} ({q1:.{digits}f}–{q3:.{digits}f}); N={n}"


def change_baseline_to_post_minmax(df, score_col, digits=1):
    tmp = df[["patient", "session", score_col]].copy()
    tmp[score_col] = pd.to_numeric(tmp[score_col], errors="coerce")

    # keep last non-missing per patient/session
    tmp = (tmp.sort_values(["patient", "session"])
              .dropna(subset=[score_col])
              .drop_duplicates(["patient", "session"], keep="last"))

    wide = tmp.pivot(index="patient", columns="session", values=score_col)

    if "pre" not in wide.columns or "post" not in wide.columns:
        return "–"

    change = wide["post"] - wide["pre"]
    change = change.dropna()

    n = change.notna().sum()
    if n == 0:
        return "–"

    q1, q3 = change.quantile([0.25, 0.75])

    return (
        f"{change.median():.{digits}f} "
        f"({q1:.{digits}f}–{q3:.{digits}f}) "
        f"[{change.min():.{digits}f}–{change.max():.{digits}f}]; "
        f"N={n}"
    )



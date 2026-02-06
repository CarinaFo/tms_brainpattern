# -*- coding: utf-8 -*-
"""
Table 1: demographics
Author: Carina Forster
"""

import pandas as pd
import re
from pathlib import Path
import matplotlib.pyplot as plt

from docx import Document

# ---------------------------
# Plot settings
# ---------------------------
plt.rcParams['pdf.fonttype'] = 42
plt.rcParams.update({
    "font.family": "sans-serif",
    "font.sans-serif": ["DejaVu Sans"],
    "font.size": 12,
    "axes.labelsize": 14,
    "axes.titlesize": 14,
    "xtick.labelsize": 12,
    "ytick.labelsize": 13,
    "legend.fontsize": 12,
})

N_patients = 70

save_dir = Path("C:/Users/CarinaF/tms_brainpattern/tables")

# ---------------------------
# Cleaning functions
# ---------------------------
def clean_participant_id(x):
    if pd.isna(x):
        return pd.NA
    match = re.search(r"(\d{3}R?)", str(x))
    return match.group(1) if match else pd.NA

# ---------------------------
# Load data
# ---------------------------
def load_and_clean_data():
    # Participant info
    df = pd.read_excel(
        Path("L:/Lab_LucaC/A_QNC_Databank/Participants_Clinical_TMS_Data/03022026 MDD Anonymised QNC Clinical Data.xlsx"),
        sheet_name='Participant Info'
    )

    # Clean participant IDs
    df['patient'] = df['Participant ID'].apply(clean_participant_id)

    # Filter to included patients
    included_ids = pd.read_csv(
        Path("L:/Lab_LucaC/Carina/canonical_hmm_finalsample/prepared_data_giles_05Hz_1Hzfiltereddata/patients_fitted_for_this_hmm.csv")
    )["patient_id"]

    df = df[df['patient'].isin(included_ids)].copy()

    # Drop original ID column
    df.drop('Participant ID', axis=1, inplace=True)

    # make sure we have 70 patients
    assert len(df) == N_patients

    return df, included_ids


def make_table1():

    df, included_ids = load_and_clean_data()

    # make table
    rows = []
    N_total = len(df)

    rows.append({
        "Demographic measure": "N",
        "Total": f"{N_total}"
    })

    rows.append({
        "Demographic measure": "Age, years — Median (IQR) [min-max]",
        "Total": median_iqr_minmax(df["Age"])
    })

    rows.append({"Demographic measure": "Gender", "Total": ""})

    gender_stats = n_percent(df["Gender"])
    for level, value in gender_stats.items():
        rows.append({
            "Demographic measure": f"  {level}",
            "Total": value
        })

    rows.append({"Demographic measure": "Handedness", "Total": ""})

    hand_stats = n_percent(df["Handedness"])
    for level, value in hand_stats.items():
        rows.append({
            "Demographic measure": f"  {level}",
            "Total": value
        })

    rows.append({
        "Demographic measure": "Years of education — Median (IQR) [min-max]",
        "Total": median_iqr_minmax(df["Years of Education"])
    })

    rows.append({
        "Demographic measure": "Age of diagnosis, years — Median (IQR) [min-max]",
        "Total": median_iqr_minmax(df["Age of Diagnosis"])
    })

    rows.append({
        "Demographic measure": "Age of symptom onset, years — Median (IQR) [min-max]",
        "Total": median_iqr_minmax(df["Age of Symptom Onset"])
    })

    rows.append({"Demographic measure": "Diagnostic type", "Total": ""})

    group_stats = n_percent(df["Diagnostic Type"])
    for level, value in group_stats.items():
        rows.append({
            "Demographic measure": f"  {level}",
            "Total": value
        })
    
    yes_no_block(df["Previous TMS"], "Previous rTMS")
    yes_no_block(df["Previous ECT"], "Previous ECT")

    rows.append({
        "Demographic measure": "rTMS sessions, n — Median (IQR) [min-max]",
        "Total": median_iqr_minmax(df["Number of Treatment Days"])
    })

    # Participant info
    madrs = pd.read_excel(
        Path("L:/Lab_LucaC/A_QNC_Databank/Participants_Clinical_TMS_Data/03022026 MDD Anonymised QNC Clinical Data.xlsx"),
        sheet_name='MADRS'
    )

    # Clean participant IDs
    madrs['patient'] = madrs['Participant ID'].apply(clean_participant_id)

    madrs = madrs[madrs['patient'].isin(included_ids)].copy()

    # convert to datetime
    madrs["Date Completed"] = pd.to_datetime(madrs["Date Completed"], errors="coerce")

    dates = (
        madrs
        .pivot(index="patient", columns="Session", values="Date Completed")
    )

    # wrong entry by technician
    dates.loc["213", "Post"] = pd.Timestamp("2025-11-03")

    dates["days_pre_to_post"] = (dates["Post"] - dates["Pre"]).dt.days

    assert (dates["days_pre_to_post"] >= 0).all()

    rows.append({
    "Demographic measure": "Days from pre- to post-time point — Median (IQR) [min–max]",
    "Total": median_iqr_minmax(dates["days_pre_to_post"])
    })

    table1 = pd.DataFrame(rows)

    dataframe_to_word_table(
        table1,
        f"{save_dir}/Table1_Population_Characteristics.docx",
        title="Table 1. Population characteristics"
    )


def dataframe_to_word_table(df, filename, title=None):
    doc = Document()

    if title:
        p = doc.add_paragraph(title)
        p.style = "Normal"
        p.runs[0].bold = True

    table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1])
    table.style = "Table Grid"

    # Header
    for j, col in enumerate(df.columns):
        cell = table.cell(0, j)
        cell.text = col
        cell.paragraphs[0].runs[0].bold = True

    # Body
    for i in range(df.shape[0]):
        for j in range(df.shape[1]):
            table.cell(i + 1, j).text = str(df.iat[i, j])

    doc.save(filename)

def median_iqr(series, digits=1):
    s = pd.to_numeric(series, errors="coerce").dropna()
    if s.empty:
        return "NA"
    q1 = s.quantile(0.25)
    q3 = s.quantile(0.75)
    med = s.median()
    return f"{med:.{digits}f} ({q1:.{digits}f}–{q3:.{digits}f})"

def median_iqr_minmax(series, digits=1):
    s = pd.to_numeric(series, errors="coerce").dropna()
    if s.empty:
        return "NA"
    q1, q3 = s.quantile([0.25, 0.75])
    return (
        f"{s.median():.{digits}f} "
        f"({q1:.{digits}f}–{q3:.{digits}f}) "
        f"[{int(s.min())}–{int(s.max())}]"
    )

def n_percent(series):
    s = series.dropna()
    counts = s.value_counts()
    total = len(s)
    return {
        k: f"{v} ({100*v/total:.1f}%)"
        for k, v in counts.items()
    }

def yes_no_block(series, label):
    rows.append({"Demographic measure": label, "Total": ""})
    stats = n_percent(series)
    for level in ["Yes", "No"]:
        if level in stats:
            rows.append({
                "Demographic measure": f"  {level}",
                "Total": stats[level]
            })